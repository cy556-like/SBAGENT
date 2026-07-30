"""Export a SBAGENT scanned-PDF OCR cache (final or partial) to editable DOCX.

Usage:
    python tools/export_pdf_ocr_cache_to_docx.py "丰田生产方式-最新版.pdf"
"""

from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path


def _find_cache(documents_dir: Path, filename: str) -> tuple[Path, bool]:
    """Return cache path and whether it is a complete cache."""
    final_name = f"{filename}.sbagent-text.json"
    partial_name = f"{filename}.sbagent-text.partial.json"
    finals = sorted(documents_dir.rglob(final_name))
    if finals:
        return finals[0], True
    partials = sorted(documents_dir.rglob(partial_name))
    if partials:
        return partials[0], False
    raise FileNotFoundError(
        f"没有找到 {filename} 的OCR缓存。请先完成至少10页OCR，"
        "或确认PDF已经上传到知识库。"
    )


def _set_run_font(run, name: str = "Microsoft YaHei") -> None:
    from docx.oxml.ns import qn

    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def export_cache(filename: str, output_path: Path | None = None) -> Path:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt

    project_root = Path(__file__).resolve().parents[1]
    sys.path.insert(0, str(project_root))
    from app.config import settings
    from app.rag.document import _get_export_dir

    documents_dir = Path(settings.DOCUMENTS_DIR)
    cache_path, complete = _find_cache(documents_dir, filename)
    payload = json.loads(cache_path.read_text(encoding="utf-8"))
    pages = [page for page in payload.get("pages", []) if str(page.get("content") or "").strip()]
    if not pages:
        raise ValueError("OCR缓存中尚无可导出的文字页面。")
    pages.sort(key=lambda item: int(item.get("page", 0)))

    if output_path is None:
        # 使用应用原本的导出目录，方便现有下载接口统一读取。
        export_dir = Path(_get_export_dir())
        suffix = "文字版" if complete else "文字版-处理中"
        output_path = export_dir / f"{Path(filename).stem}-{suffix}.docx"
    output_path.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    normal = document.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia",
        "Microsoft YaHei",
    )

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(f"{Path(filename).stem} - OCR文字版")
    title_run.bold = True
    title_run.font.size = Pt(18)
    _set_run_font(title_run)

    status = "完整OCR缓存" if complete else "处理中断点缓存（可稍后重新导出完整版本）"
    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note_run = note.add_run(f"来源：{filename}  |  {status}  |  已导出 {len(pages)} 页")
    note_run.font.size = Pt(9)
    _set_run_font(note_run)

    for index, page in enumerate(pages):
        page_number = int(page.get("page_number", int(page.get("page", index)) + 1))
        heading = document.add_paragraph()
        heading.paragraph_format.space_before = Pt(14)
        heading.paragraph_format.space_after = Pt(6)
        heading_run = heading.add_run(f"原PDF第 {page_number} 页")
        heading_run.bold = True
        heading_run.font.size = Pt(13)
        _set_run_font(heading_run)

        text = str(page.get("content") or "").strip()
        for line in text.splitlines():
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(3)
            paragraph.paragraph_format.line_spacing = 1.35
            run = paragraph.add_run(line.strip())
            _set_run_font(run)

        if index + 1 < len(pages):
            document.add_page_break()

    document.save(output_path)
    return output_path


def main() -> int:
    parser = argparse.ArgumentParser(description="将SBAGENT的PDF OCR缓存导出为可编辑Word")
    parser.add_argument("filename", help="知识库中的PDF文件名，例如：丰田生产方式-最新版.pdf")
    parser.add_argument("--out", help="输出docx路径（可选）")
    args = parser.parse_args()
    output = export_cache(args.filename, Path(args.out) if args.out else None)
    print(f"WORD_EXPORT_OK: {output}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
